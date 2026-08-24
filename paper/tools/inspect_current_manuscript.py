"""Print the structural outline and figure/table captions of the current manuscript."""

from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
MANUSCRIPT = ROOT / "paper" / "newest_original_manuscript_llm_benchmark_scope_corrected.docx"


def main() -> None:
    document = Document(MANUSCRIPT)
    print(f"paragraphs={len(document.paragraphs)} tables={len(document.tables)} inline_shapes={len(document.inline_shapes)}")
    for index, paragraph in enumerate(document.paragraphs):
        text = " ".join(paragraph.text.split())
        style = paragraph.style.name if paragraph.style else ""
        if (
            text
            and (
                95 <= index <= 194
                or
                style.startswith("Heading")
                or text.startswith("Figure")
                or text.startswith("Table")
                or "SHAP" in text
                or "LIME" in text
                or "observed" in text.lower()
                or "predicted" in text.lower()
            )
        ):
            print(f"P{index:04d}\t{style}\t{text}")


if __name__ == "__main__":
    main()
