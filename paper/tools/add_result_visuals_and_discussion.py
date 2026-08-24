"""Insert prediction and XAI result visuals into the current manuscript."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "paper" / "newest_original_manuscript_llm_benchmark_scope_corrected.docx"
OUTPUT = ROOT / "paper" / "newest_original_manuscript_results_visuals_integrated.docx"
FIGURE_8 = (
    ROOT
    / "outputs"
    / "final_five_model_prediction_visuals_v1"
    / "observed_vs_predicted_scatter_oos_2022_2025.png"
)
FIGURE_9 = (
    ROOT
    / "outputs"
    / "final_five_model_prediction_visuals_v1"
    / "actual_vs_predicted_oos_2025_zoom.png"
)
FIGURE_7 = ROOT / "paper" / "assets" / "figure10_shap_lime_result_audit.png"

FIGURE_ALT_TEXT = [
    "Five-pillar point-in-time reliability-audit pipeline for SET50 forecasting.",
    "Point-in-time financial-news normalization, filtering, sentiment prediction and session aggregation.",
    "Expanding-window evaluation with purged next-day labels and held-out outer test years.",
    "Numerical feature construction, causal rolling VMD and the five registered neural architectures.",
    "Separate downstream SET50 news-forecasting audit and locked intrinsic Bull-Bear-Leader sentiment audit.",
    "Causal Bull-Sideway-Bear routing, training-only SHAP selection and paired temporal inference.",
    "Regime-SHAP balanced-accuracy effects and model-by-regime LIME low-fidelity rates.",
    "Observed-versus-predicted next-day SET50 scatter plots for the five frozen architectures.",
    "Actual and predicted next-day SET50 level series for the five architectures in the 2025 fold.",
    "Heatmap of architecture-wise audit effects and intrinsic Leader sentiment gains.",
]


def set_run_font(run, *, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Times New Roman")
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def copy_paragraph_properties(source, target) -> None:
    source_properties = source._p.pPr
    if source_properties is None:
        return
    target_properties = target._p.pPr
    if target_properties is not None:
        target._p.remove(target_properties)
    target._p.insert(0, deepcopy(source_properties))


def add_body_before(anchor, text: str, body_template):
    paragraph = anchor.insert_paragraph_before()
    copy_paragraph_properties(body_template, paragraph)
    set_run_font(paragraph.add_run(text))
    return paragraph


def add_heading_before(anchor, text: str, heading_template):
    paragraph = anchor.insert_paragraph_before()
    copy_paragraph_properties(heading_template, paragraph)
    set_run_font(paragraph.add_run(text), bold=True)
    return paragraph


def add_figure_before(anchor, image_path: Path, *, width_inches: float, page_break: bool = False):
    paragraph = anchor.insert_paragraph_before()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = page_break
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    return paragraph


def add_caption_before(anchor, number: int, text: str, caption_template):
    paragraph = anchor.insert_paragraph_before()
    copy_paragraph_properties(caption_template, paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    label = paragraph.add_run(f"Figure {number}.")
    set_run_font(label, bold=True)
    set_run_font(paragraph.add_run(f" {text}"))
    return paragraph


def find_paragraph(document: Document, exact_text: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


def find_paragraph_start(document: Document, prefix: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph prefix not found: {prefix}")


def move_paragraph_before(paragraph, anchor) -> None:
    """Move an existing paragraph immediately before an anchor paragraph."""

    anchor._p.addprevious(paragraph._p)


def replace_paragraph_prefix(paragraph, old: str, new: str) -> None:
    text = paragraph.text
    if not text.startswith(old):
        raise ValueError(f"Expected paragraph prefix {old!r}: {text!r}")
    paragraph.text = new + text[len(old) :]
    for run in paragraph.runs:
        set_run_font(run)


def renumber_figure_caption(paragraph, old_number: int, new_number: int) -> None:
    prefix = f"Figure {old_number}."
    text = paragraph.text.strip()
    if not text.startswith(prefix):
        raise ValueError(f"Expected caption prefix {prefix!r}: {text!r}")
    caption_text = text[len(prefix) :].strip()
    paragraph.clear()
    label = paragraph.add_run(f"Figure {new_number}.")
    set_run_font(label, bold=True)
    set_run_font(paragraph.add_run(f" {caption_text}"))


def set_figure_alt_text(document: Document) -> None:
    drawing_properties = document._element.xpath(".//wp:docPr")
    if len(drawing_properties) != len(FIGURE_ALT_TEXT):
        raise ValueError(
            f"Expected {len(FIGURE_ALT_TEXT)} inline figures, found {len(drawing_properties)}"
        )
    for number, (properties, description) in enumerate(
        zip(drawing_properties, FIGURE_ALT_TEXT, strict=True), start=1
    ):
        properties.set("name", f"Figure {number}")
        properties.set("title", f"Figure {number}")
        properties.set("descr", description)


def main() -> None:
    for path in (SOURCE, FIGURE_7, FIGURE_8, FIGURE_9):
        if not path.exists():
            raise FileNotFoundError(path)

    document = Document(SOURCE)
    discussion_heading = find_paragraph(document, "5. Discussion")
    section_45_heading = find_paragraph_start(document, "4.5 Partial")
    body_template = find_paragraph_start(document, "Taken together, the results indicate")
    caption_template = find_paragraph_start(document, "Figure 6.")

    # Keep the SHAP/LIME evidence beside Table 4 and its fidelity result.
    shap_discussion = find_paragraph_start(document, "The regime result should also be interpreted")
    add_body_before(
        shap_discussion,
        "Figure 7 places the selector effect beside the independent local-fidelity diagnostic. Regime-SHAP "
        "improved CNN by 1.46 percentage points and LSTM-CNN by 0.05 points, while changing LSTM, "
        "LSTM-Attention and LSTM-CNN-Attention by -0.10, -1.03 and -0.74 points, respectively. Every 95% "
        "interval crossed zero and no contrast survived Holm adjustment. LIME classified 65.8-77.5% of "
        "the 120 repeats in each model-regime cell as low fidelity; overall, 1,293 of 1,800 explanations "
        "(71.83%) had local-surrogate R^2 below 0.70.",
        body_template,
    )
    add_figure_before(shap_discussion, FIGURE_7, width_inches=6.45)
    add_caption_before(
        shap_discussion,
        7,
        "Explainability audit. Panel A reports the Regime-SHAP minus Regime-All balanced-accuracy change "
        "with descriptive 95% intervals after seed aggregation within outer years; no effect survived Holm "
        "adjustment. Panel B reports the percentage of LIME repeats with local-surrogate R^2 below 0.70; "
        "failed repeats remain in the denominator.",
        caption_template,
    )

    # The integrated prediction diagnostics follow the paragraph that defines
    # their exact post-hoc Regime-SHAP-Numeric-News arm.
    integrated_paragraph = find_paragraph_start(
        document, "Separately, the post-hoc integrated 2×2 robustness analysis"
    )
    move_paragraph_before(integrated_paragraph, section_45_heading)
    add_body_before(
        section_45_heading,
        "Figures 8 and 9 visualize the five seed-averaged out-of-sample prediction series from the "
        "post-hoc Regime-SHAP-Numeric-News arm. Over 2022-2025, LSTM produced the smallest level errors "
        "(RMSE 15.97; MAE 12.73) and the tightest scatter around the identity line, but its mean balanced "
        "accuracy was only 52.01%. LSTM-CNN-Attention produced the highest balanced accuracy (53.64%) "
        "and MCC (0.089) but the largest level errors (RMSE 31.77; MAE 25.34). LSTM-Attention occupied the "
        "middle ground (BAcc 52.62%; RMSE 20.41), LSTM-CNN improved directional performance (52.81%) at "
        "the cost of larger level error (RMSE 29.79), and CNN had the weakest mean directional result "
        "(51.49%) despite visible level co-movement (RMSE 22.74). Thus, the model ranking depends on whether "
        "the endpoint is next-day level error or the sign of the next-day move.",
        body_template,
    )
    add_figure_before(section_45_heading, FIGURE_8, width_inches=6.45)
    add_caption_before(
        section_45_heading,
        8,
        "Observed versus predicted next-day SET50 levels for the post-hoc integrated arm, 2022-2025. Each "
        "point is a seed-averaged out-of-sample forecast; the dashed identity line indicates exact level "
        "agreement. RMSE and MAE evaluate level error, whereas balanced accuracy evaluates direction.",
        caption_template,
    )
    add_figure_before(section_45_heading, FIGURE_9, width_inches=6.45, page_break=True)
    add_caption_before(
        section_45_heading,
        9,
        "Seed-averaged out-of-sample next-day SET50 forecasts in the 2025 fold. Black lines show realized "
        "next-day closes and colored lines show model forecasts; annotations report the metrics calculated "
        "only on the displayed fold. The visual diagnoses level tracking and turning-point behavior, while "
        "balanced accuracy remains the primary directional endpoint.",
        caption_template,
    )
    add_body_before(
        section_45_heading,
        "In the 2025 fold, LSTM-CNN achieved the best balanced accuracy (55.15%) and direction accuracy "
        "(54.27%) but an RMSE of 30.41, whereas LSTM retained the lowest RMSE (21.57) with balanced "
        "accuracy of 53.20%. LSTM-Attention reached 53.82% balanced accuracy; CNN was near chance at "
        "51.07%; and LSTM-CNN-Attention fell to 53.07% while its RMSE rose to 36.12. The realized Up share "
        "was 48.72%, yet the models predicted Up on 74.79-86.32% of dates; corresponding Down recall was "
        "only 16.67-28.33%. The broad level paths can therefore appear visually convincing even when local "
        "reversals and small negative moves are missed.",
        body_template,
    )
    add_body_before(
        section_45_heading,
        "Taken model by model, LSTM was strongest at level calibration but remained directionally "
        "asymmetric; CNN captured broad level co-movement but smoothed or plateaued around turning points "
        "and provided little net sign discrimination; LSTM-CNN gained 2025 directional sensitivity but "
        "produced more dispersed levels; LSTM-Attention offered a compromise with smaller level error than "
        "the convolutional hybrids but remained Up-skewed; and LSTM-CNN-Attention had the best four-year "
        "primary metric while showing the largest 2025 error and weakest Down recall. This pattern is "
        "consistent with added convolution and attention capacity trading level calibration for sensitivity "
        "to selected temporal patterns, but the audit cannot identify that mechanism causally. CNN's "
        "positive Regime-SHAP cell is compatible with feature compression helping a local-filter "
        "architecture, whereas LIME's low fidelity prevents a stable feature-level explanation from being "
        "claimed (Lundberg & Lee, 2017; Ribeiro et al., 2016; Yeo et al., 2025).",
        body_template,
    )

    # The cross-pillar summary remains after all pillar-specific results and is
    # therefore renumbered from Figure 7 to Figure 10.
    summary_intro = find_paragraph_start(document, "Figure 7 summarizes architecture-wise")
    replace_paragraph_prefix(summary_intro, "Figure 7", "Figure 10")
    summary_caption = find_paragraph_start(document, "Figure 7. Architecture-wise")
    renumber_figure_caption(summary_caption, 7, 10)

    intrinsic_discussion = find_paragraph_start(document, "The intrinsic LLM experiment addresses")
    add_body_before(
        intrinsic_discussion,
        "The prediction-level diagnostics further clarify why high visual agreement should not be equated "
        "with a useful direction classifier. Across 2022-2025, observed-predicted level correlations ranged "
        "from 0.927 to 0.983, yet balanced accuracy spanned only 51.49-53.64%. Because the next-day index "
        "level inherits substantial persistence from the current level, a model can track the level closely "
        "while placing the predicted change on the wrong side of zero; LSTM's low RMSE and modest balanced "
        "accuracy illustrate this distinction. Conversely, hybrid models improved sign classification in "
        "selected folds while sacrificing calibration and Down-class recall. This endpoint mismatch is "
        "consistent with the distinction between continuous-price forecasting and direction classification "
        "in financial deep learning (Fischer & Krauss, 2018; Hoseinzadeh & Haratizadeh, 2019; Sezer et al., "
        "2020). The SHAP-LIME panel adds a second caution: a positive selector ablation for one architecture "
        "does not establish that its individual attributions are locally faithful or transportable.",
        body_template,
    )

    set_figure_alt_text(document)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
